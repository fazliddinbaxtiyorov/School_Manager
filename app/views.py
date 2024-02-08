from django.shortcuts import render, redirect
from django.http import FileResponse
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
from io import BytesIO
# Create your views here.
from .forms import NewsForm, AchievementForm, TeachersForm, StudentsForm, InformationSchoolForm
from .models import News, StudentsAchievement, Teachers, Students, InformationSchool


def home(request):
    teachers = Teachers.objects.all()
    st_achievement = StudentsAchievement.objects.all()
    for_news = News.objects.all()
    return render(request, 'index.html', {'news': for_news, 'teachers': teachers, 'achivement': st_achievement})


def admin(request):
    student = Students.objects.all()
    return render(request, 'admin.html', {'student': student})


def xizmat(request):
    return render(request, 'xizmatlar.html')


def hi(request):
    return render(request, 'hi.html')


def contact(request):
    return render(request, 'contact.html')


def malumotnoma(request):
    return render(request, 'malumotnoma.html')


def adding_students(request):
    if request.method == 'POST':
        form = StudentsForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('hi')
    else:
        form = StudentsForm()
    return render(request, 'add.html', {'form': form})


def adding_news(request):
    if request.method == 'POST':
        form = NewsForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('/')
    else:
        form = NewsForm()
    return render(request, 'add.html', {'form': form})


def adding_achievement(request):
    if request.method == 'POST':
        form = AchievementForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('/')
    else:
        form = AchievementForm()
    return render(request, 'add.html', {'form': form})


def adding_teachers(request):
    if request.method == 'POST':
        form = TeachersForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('/')
    else:
        form = TeachersForm()
    return render(request, 'add.html', {'form': form})


def adding(request):
    if request.method == 'POST':
        form = InformationSchoolForm(request.POST, request.FILES)
        if form.is_valid():
            name = form.cleaned_data.get('title')
            grade = form.cleaned_data.get('grade')
            request.session['name'] = name
            request.session['grade'] = grade
            form.save()
            return redirect('doc')
    else:
        form = InformationSchoolForm()
    return render(request, 'add.html', {'form': form})


def doc_school(request):
    doc = Document()
    doc.add_heading("Ma'lumotnoma", level=3)

    name = request.session.get('name')
    grade = request.session.get('grade')

    if name is not None and grade is not None:
        today = date.today()
        doc.add_paragraph(
            f"Olmaliq shaxar 18-umumiy o'rta ta'lim maktabining {grade}-sinfida {name} haqiqatdan ham taxsil oladi.")
        doc.add_picture("/Users/acer/OneDrive/Рабочий стол/School-Manager/app/pechat.jpg", width=Inches(1.3),
                        height=Inches(1.4))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.add_paragraph(f"{today}")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = FileResponse(buffer, as_attachment=True, filename='information.docx')
        return response
    else:
        return HttpResponse("Name and grade values not found in session.")